from docx import Document
from docx.shared import Pt
from io import BytesIO
from datetime import datetime


def create_docx(report):

    document = Document()

    title = document.add_heading("AI Meeting Notes Report", level=1)
    title.runs[0].font.size = Pt(20)

    document.add_paragraph(
        f"Generated on: {datetime.now().strftime('%d %B %Y %H:%M')}"
    )

    document.add_paragraph()

    current_heading = None

    for line in report.split("\n"):

        line = line.strip()

        if not line:
            continue

        if line.startswith("#"):

            current_heading = line.replace("#", "").strip()

            document.add_heading(current_heading, level=2)

        elif line.startswith("•"):

            document.add_paragraph(
                line.replace("•", "").strip(),
                style="List Bullet"
            )

        else:

            line = (
                line.replace("**", "")
                    .replace("__", "")
                    .replace("*", "")
                    .replace("`", "")
            )

            document.add_paragraph(line)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer